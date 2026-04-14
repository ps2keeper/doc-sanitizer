"""Word document (.docx) processing handler with track-changes support."""
import re
from copy import deepcopy
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from engine.audit_engine import AuditEngine, AuditResult


class DocxHandler:
    """Handles .docx file processing: replaces sensitive words with track-changes."""

    def __init__(self):
        self._revision_id = 0

    def _next_id(self) -> str:
        self._revision_id += 1
        return str(self._revision_id)

    def process(self, file_path: str, replacements: dict[str, str]) -> tuple[BytesIO, AuditResult, dict[str, int]]:
        """Process a Word document by replacing sensitive words with tracked deletions/insertions."""
        self._revision_id = 0
        doc = Document(file_path)

        # Count matches before replacing
        replacement_counts = {}
        for word, replacement in replacements.items():
            count = 0
            for p in doc.paragraphs:
                count += len(re.findall(re.escape(word), p.text, re.IGNORECASE))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            count += len(re.findall(re.escape(word), p.text, re.IGNORECASE))
            if count > 0:
                replacement_counts[word] = count

        # Enable track changes
        self._enable_track_changes(doc)

        # Process all paragraphs (including table cells)
        self._process_document(doc, replacements)

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Audit: extract all text and scan
        full_text = self._extract_text(doc)
        audit_engine = AuditEngine(replacements)
        audit_result = audit_engine.scan(full_text)

        return output, audit_result, replacement_counts

    def _enable_track_changes(self, doc):
        """Enable track changes in the document settings."""
        settings = doc.settings.element
        track_el = settings.find(qn('w:trackRevisions'))
        if track_el is None:
            new_settings = Document().settings.element
            track_el = new_settings.makeelement(qn('w:trackRevisions'), {})
            settings.append(track_el)

    def _process_document(self, doc, replacements):
        """Process all paragraphs and tables in the document."""
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, replacements)

    def _process_paragraph(self, paragraph, replacements):
        """Process a single paragraph in-place, preserving all paragraph properties and run formatting."""
        full_text = paragraph.text
        if not full_text:
            return

        # Find all sensitive word positions in the full text
        all_matches = []
        for word, replacement in replacements.items():
            for match in re.finditer(re.escape(word), full_text, re.IGNORECASE):
                all_matches.append((match.start(), match.end(), match.group(), replacement))

        if not all_matches:
            return

        all_matches.sort(key=lambda x: x[0])

        # Get original runs in order
        runs = []
        for child in list(paragraph._p):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                runs.append(child)

        # Build position-to-run mapping
        # We need to know which run contains which text position
        run_positions = []  # (start_pos, end_pos, run_element)
        pos = 0
        for run_el in runs:
            # Find all w:t elements in this run
            texts = []
            for child in run_el:
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if ctag == 't' and child.text:
                    texts.append(child.text)
            text_content = ''.join(texts)
            run_positions.append((pos, pos + len(text_content), run_el))
            pos += len(text_content)

        # For each match, find which run(s) it spans and process in-place
        # Process from RIGHT to LEFT so earlier replacements don't shift later positions
        for match_start, match_end, original, replacement in reversed(all_matches):
            # Find the run that contains this match
            target_run = None
            target_run_start = 0
            for rp_start, rp_end, rp_el in run_positions:
                if match_start >= rp_start and match_end <= rp_end:
                    target_run = rp_el
                    target_run_start = rp_start
                    break

            if target_run is None:
                continue

            # Get all w:t elements in this run
            t_elements = []
            for child in target_run:
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if ctag == 't' and child.text:
                    t_elements.append(child)

            if not t_elements:
                continue

            # Calculate local position within the combined text of all t elements
            combined = ''.join(t.text for t in t_elements)
            local_start = match_start - target_run_start
            local_end = match_end - target_run_start

            if local_start < 0 or local_end > len(combined):
                continue

            # Find which t element(s) the match spans
            t_positions = []
            tpos = 0
            for t_el in t_elements:
                t_positions.append((tpos, tpos + len(t_el.text), t_el))
                tpos += len(t_el.text)

            # For simplicity, we handle the common case: match is within a single t element
            # Find the t element containing the match
            for t_start, t_end, t_el in t_positions:
                if local_start >= t_start and local_end <= t_end:
                    # Match is entirely within this single t element
                    before = t_el.text[:local_start - t_start]
                    after = t_el.text[local_end - t_start:]

                    # Replace this t element's text with before + deleted + inserted + after
                    # We need to restructure the run element
                    self._replace_in_run(target_run, t_el, before, original, replacement, after)
                    break
            else:
                # Match spans multiple t elements - handle by merging and splitting
                self._replace_multi_run(target_run, t_positions, local_start, local_end, original, replacement)

    def _replace_in_run(self, run_el, t_el, before: str, original: str, replacement: str, after: str):
        """Replace text within a single run, creating del/ins markings while preserving all run properties."""
        # Get the run's rPr (format)
        rPr = run_el.find(qn('w:rPr'))

        # Collect all t elements from this run
        t_elements = []
        for child in list(run_el):
            ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if ctag == 't':
                t_elements.append(child)

        # Helper to create a w:t element
        def make_t(text_content):
            t = run_el.makeelement(qn('w:t'), {})
            t.text = text_content
            for child in t_elements:
                if child.get(qn('xml:space')):
                    t.set(qn('xml:space'), 'preserve')
                    break
            return t

        # Helper to create a full w:r element with the run's format + optional extras
        def make_r(text, extras):
            new_r = run_el.makeelement(qn('w:r'), {})
            if rPr is not None:
                new_rPr = deepcopy(rPr)
            else:
                new_rPr = run_el.makeelement(qn('w:rPr'), {})
            new_r.append(new_rPr)
            # Add extra markings (strike, del, highlight, ins)
            for extra in extras:
                new_rPr.append(extra)
            new_r.append(make_t(text))
            return new_r

        # Helper to get paragraph element
        parent = run_el.getparent()

        # Remove all w:t elements from the original run
        for t_el_orig in t_elements:
            run_el.remove(t_el_orig)

        # Rebuild the original run with before text (if any)
        if before:
            run_el.append(make_t(before))
        else:
            # Keep at least one empty t element so the run isn't invalid
            run_el.append(make_t(''))

        # Create deleted run (insert AFTER the original run in paragraph)
        strike = run_el.makeelement(qn('w:strike'), {qn('w:val'): 'true'})
        del_mark = run_el.makeelement(qn('w:del'), {
            qn('w:id'): self._next_id(),
            qn('w:author'): 'DocSanitizer',
            qn('w:date'): '2026-04-14T00:00:00Z'
        })
        del_r = make_r(original, [strike, del_mark])
        run_el.addnext(del_r)

        # Create inserted run (insert AFTER the deleted run)
        highlight = run_el.makeelement(qn('w:highlight'), {qn('w:val'): 'yellow'})
        ins_mark = run_el.makeelement(qn('w:ins'), {
            qn('w:id'): self._next_id(),
            qn('w:author'): 'DocSanitizer',
            qn('w:date'): '2026-04-14T00:00:00Z'
        })
        ins_r = make_r(replacement, [highlight, ins_mark])
        del_r.addnext(ins_r)

        # Handle after text: if there's after text, keep it in the original run
        # Actually, the original run now has before text. We need a new run for after text.
        if after:
            after_r = run_el.makeelement(qn('w:r'), {})
            if rPr is not None:
                after_rPr = deepcopy(rPr)
            else:
                after_rPr = run_el.makeelement(qn('w:rPr'), {})
            after_r.append(after_rPr)
            after_r.append(make_t(after))
            ins_r.addnext(after_r)

    def _replace_multi_run(self, run_el, t_positions, local_start, local_end, original, replacement):
        """Handle a match that spans multiple w:t elements."""
        combined = ''.join(t.text for t_start, t_end, t_el in t_positions)
        rPr = run_el.find(qn('w:rPr'))

        # Remove all old t elements
        for _, _, t_el in t_positions:
            run_el.remove(t_el)

        def make_t(text_content):
            t = run_el.makeelement(qn('w:t'), {})
            t.text = text_content
            if any(t_el.get(qn('xml:space')) for _, _, t_el in t_positions):
                t.set(qn('xml:space'), 'preserve')
            return t

        def make_r(text, extras):
            new_r = run_el.makeelement(qn('w:r'), {})
            if rPr is not None:
                new_rPr = deepcopy(rPr)
            else:
                new_rPr = run_el.makeelement(qn('w:rPr'), {})
            new_r.append(new_rPr)
            for extra in extras:
                new_rPr.append(extra)
            new_r.append(make_t(text))
            return new_r

        before = combined[:local_start]
        after = combined[local_end:]

        if before:
            run_el.append(make_t(before))
        else:
            run_el.append(make_t(''))

        if rPr is None:
            rPr = run_el.makeelement(qn('w:rPr'), {})
            run_el.insert(0, rPr)

        strike = run_el.makeelement(qn('w:strike'), {qn('w:val'): 'true'})
        del_mark = run_el.makeelement(qn('w:del'), {
            qn('w:id'): self._next_id(),
            qn('w:author'): 'DocSanitizer',
            qn('w:date'): '2026-04-14T00:00:00Z'
        })
        del_r = make_r(original, [strike, del_mark])
        run_el.addnext(del_r)

        highlight = run_el.makeelement(qn('w:highlight'), {qn('w:val'): 'yellow'})
        ins_mark = run_el.makeelement(qn('w:ins'), {
            qn('w:id'): self._next_id(),
            qn('w:author'): 'DocSanitizer',
            qn('w:date'): '2026-04-14T00:00:00Z'
        })
        ins_r = make_r(replacement, [highlight, ins_mark])
        del_r.addnext(ins_r)

        if after:
            after_r = run_el.makeelement(qn('w:r'), {})
            after_rPr = deepcopy(rPr)
            after_r.append(after_rPr)
            after_r.append(make_t(after))
            ins_r.addnext(after_r)

    def _extract_text(self, doc) -> str:
        """Extract all text from a document, skipping deleted (w:del) text."""
        parts = []
        for p in doc.paragraphs:
            parts.append(self._extract_paragraph_text(p))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(self._extract_paragraph_text(p))
        return '\n'.join(parts)

    def _extract_paragraph_text(self, paragraph) -> str:
        """Extract text from a paragraph, skipping text inside deleted runs."""
        texts = []
        for child in paragraph._p:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                # Check if this is a normal run or deleted run
                rPr = child.find(qn('w:rPr'))
                if rPr is not None and rPr.find(qn('w:del')) is not None:
                    continue
                # Check if this is an inserted run (include it)
                if rPr is not None and rPr.find(qn('w:ins')) is not None:
                    for t in child:
                        if t.tag.split('}')[-1] == 't' and t.text:
                            texts.append(t.text)
                    continue
                # Normal run
                for t in child:
                    if t.tag.split('}')[-1] == 't' and t.text:
                        texts.append(t.text)
        return ''.join(texts)
