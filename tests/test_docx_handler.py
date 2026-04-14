import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from engine.docx_handler import DocxHandler


def test_basic_replacement_track_changes(tmp_path):
    from docx import Document
    # Create test document
    doc = Document()
    doc.add_paragraph("This is secret information.")
    doc_path = str(tmp_path / "test.docx")
    doc.save(doc_path)

    handler = DocxHandler()
    result, audit = handler.process(doc_path, {'secret': 'REDACTED'})
    processed_content = result.getvalue()

    # Result should be a valid .docx
    assert processed_content is not None
    assert len(processed_content) > 0

    # Audit should confirm no "secret" remains
    assert audit.is_clean is True


def test_replacement_in_table(tmp_path):
    from docx import Document
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "confidential data"
    doc_path = str(tmp_path / "test.docx")
    doc.save(doc_path)

    handler = DocxHandler()
    result, audit = handler.process(doc_path, {'confidential': '[CONF]', 'data': '[DATA]'})
    assert audit.is_clean is True


def test_no_sensitive_words(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a clean document.")
    doc_path = str(tmp_path / "test.docx")
    doc.save(doc_path)

    handler = DocxHandler()
    result, audit = handler.process(doc_path, {})
    assert audit.is_clean is True
